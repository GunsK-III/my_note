# 2024/09/20
# 在Word文档中用字符画画 (●'◡'●)
"""相比于在excel中使用单元格作画，用字符作画显示效果较差。唯有在显示人的大头照时，有不错的效果。"""

'''第一步：压缩图片尺寸。'''
import cv2

print("————————————————————————————————————————————————————————————————————————————————\n"
      "与上一个excel画图项目不同，这个word画图项目需要你手动创建word文件(或者下载我的“56×60.docx”word模板)。\n"
      "——————\033[91m再次提醒：你准备的图片的长宽比例，要和word中字符排版的长宽比例尽量一样。\033[0m——————\n"
      "我上传的那个word模板为a4纸，显示方向是竖向，页边距、行距已调为最小，整页文字长宽比约为1比1.4，\n"
      "文字为笔画较多的“饕餮”二字，列数为56，行数为60。为了显示效果更好，它的排版与正常word差异较大。\n"
      "或者，你可以为了让显示效果更好一点，在word中设置字体加粗、显示方向为横向、文字沾满第一页等等……\n"
      "最后这个word可能已经改的不像样了，但是写入图片的像素值后显示效果会不错。 (～￣▽￣)～\n"
      "再加一点：为了显示效果更好，可以在Windows图片应用中进行编辑，拉高图片对比度、突出显示等。")

image_path = input(r"输入图片的路径(路径前后不要用引号): ")
img = cv2.imread(image_path)
if img is None:
    print("路径输入有误！")

width, height = img.shape[1], img.shape[0]
print("图片的宽度为：", width, "图片的高度为：", height)
print("下面分别输入图片压缩后的宽度和高度。"
      "为了最后显示效果美观，建议了解图片原尺寸后输入等比例的宽高数值")

new_width = int(input("输入文档第一页字符的列数(正整数)："))
new_height = int(input("输入文档第一页字符的行数(正整数)："))
resized = cv2.resize(img, (new_width, new_height))
if_write = input("如果要保存压缩后的图片，输入“是”或“y”，否则键入“enter”：")
if if_write == "y" or if_write == "是":
    print("图片默认保存在与本文件相同路径下。")
    output_path = 'yasuohou.jpg'
    cv2.imwrite(output_path, resized)
    print("压缩图片已保存！")
else:
    pass

'''第二步：将图片中像素的颜色赋给word中的字体。'''
import cv2
from docx import Document
from docx.shared import RGBColor

doc_path = input(r"输入word文档的路径(路径前后不要用引号): ")
doc = Document(doc_path)
resized = cv2.cvtColor(resized, cv2.COLOR_BGR2RGB)
rows = new_height
cols = new_width

if resized.shape[0] != rows or resized.shape[1] != cols:  # 检查图片尺寸是否与文字尺寸匹配
    raise ValueError("图片尺寸与文字行列数不匹配。")

first_page_text = doc.paragraphs[0].text  # 所有文字都应在第一个段落中
total_chars = len(first_page_text)

fst_pag = doc.paragraphs[0]
fst_run = fst_pag.runs[0]

# 开始遍历文档中的每个文字，并设置颜色
for i in range(total_chars):
    row = i // cols
    col = i % cols
    pixel_color = resized[row, col]
    # 我想了想，决定不改变原来字体的颜色，而是写入新的带有颜色的字体。以免在显示效果不好的情况下还把原稿破坏了。
    run = doc.paragraphs[0].add_run(first_page_text[i])
    run.font.color.rgb = RGBColor(int(pixel_color[0]), int(pixel_color[1]), int(pixel_color[2]))
    run.font.size = fst_run.font.size  # 除了修改颜色，还要保证字体格式不变。
    run.font.name = fst_run.font.name
    run.font.bold = fst_run.font.bold

doc.save(r"D:\NewFolder\pythonProject1\.venv\Files_Crate\PyDesign\79×42.docx")
print("彩色字体已添加到word中。\n"
      "离远点看效果会更好。(❁´◡`❁)\n")
input("程序已结束。按回车键退出。")
